# backend/app/services/document_service.py
"""
Improved Document Service with AI integration
"""

import os
import uuid
from typing import Dict, Any, Optional, List
from datetime import datetime
import PyPDF2
import docx
from io import BytesIO
from pathlib import Path
import logging

from app.core.config import settings
from app.services.ai_service import ai_service
from app.database import get_database
from fastapi import HTTPException, UploadFile

logger = logging.getLogger(__name__)


class DocumentService:
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)
        
        # File validation settings
        self.allowed_extensions = {'.pdf', '.docx', '.doc', '.txt'}
        self.max_file_size = settings.MAX_FILE_SIZE
        
    async def process_cv_upload(self, file: UploadFile, user_id: str) -> Dict[str, Any]:
        """
        Process CV upload with AI-powered parsing
        """
        try:
            # Validate file
            self._validate_file(file)
            
            # Read file content
            await file.seek(0)
            file_content = await file.read()
            
            # Save file
            file_path = await self._save_file(file_content, file.filename, user_id)
            
            # Extract text from file
            text_content = self._extract_text(file_content, Path(file.filename).suffix.lower())
            
            # Process with AI
            ai_result = await ai_service.process_cv_with_ai(text_content, user_id)
            
            # Create document record
            document_record = {
                "user_id": user_id,
                "file_info": {
                    "original_filename": file.filename,
                    "file_path": str(file_path),
                    "file_size": len(file_content),
                    "content_type": file.content_type,
                    "upload_date": datetime.utcnow()
                },
                "text_content": text_content[:2000],  # Store first 2000 chars for preview
                "cv_data": ai_result.get("cv_data", {}),
                "processing_metadata": ai_result.get("metadata", {}),
                "status": ai_result.get("status", "completed"),
                "document_type": "cv",
                "is_active": True,
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow()
            }
            
            # Save to database
            db = await get_database()
            result = await db.documents.insert_one(document_record)
            
            # Sync cv_data to user profile for quick access by other services
            await db.users.update_one(
                {"_id": user_id if isinstance(user_id, object) else self._to_object_id(user_id)},
                {"$set": {
                    "cv_uploaded_at": datetime.utcnow(),
                    "cv_data": document_record["cv_data"],  # Sync to user profile
                    "active_cv_document_id": str(result.inserted_id)
                }}
            )
            
            logger.info(f"Successfully processed CV upload for user {user_id}")
            
            return {
                "success": True,
                "document_id": str(result.inserted_id),
                "file_info": document_record["file_info"],
                "cv_data": document_record["cv_data"],
                "status": document_record["status"],
                "message": "CV uploaded and processed successfully"
            }
            
        except Exception as e:
            logger.error(f"Error processing CV upload: {e}")
            # Clean up file if it was saved
            if 'file_path' in locals() and Path(file_path).exists():
                Path(file_path).unlink()
            
            raise HTTPException(
                status_code=500,
                detail=f"Failed to process CV: {str(e)}"
            )

    async def customize_cv_for_job(self, document_id: str, job_description: str, 
                                 company_name: str, user_id: str,
                                 user_preferences: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Customize CV for a specific job
        """
        try:
            # Get original CV document
            document = await self.get_document(document_id, user_id)
            if not document:
                raise HTTPException(status_code=404, detail="CV document not found")
            
            cv_data = document.get("cv_data", {})
            
            # Customize CV using AI service
            customization_result = await ai_service.customize_cv_for_job(
                cv_data=cv_data,
                job_description=job_description,
                company_name=company_name,
                user_preferences=user_preferences
            )
            
            return {
                "success": True,
                "customized_cv": customization_result.get("customized_cv", cv_data),
                "metadata": customization_result.get("metadata", {}),
                "status": customization_result.get("status", "completed")
            }
            
        except Exception as e:
            logger.error(f"Error customizing CV: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to customize CV: {str(e)}"
            )

    async def generate_cover_letter(self, document_id: str, job_description: str,
                                  company_name: str, job_title: str, user_id: str,
                                  user_preferences: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Generate cover letter for a job
        """
        try:
            # Get CV document
            document = await self.get_document(document_id, user_id)
            if not document:
                raise HTTPException(status_code=404, detail="CV document not found")
            
            cv_data = document.get("cv_data", {})
            
            # Generate cover letter using AI service
            cover_letter_result = await ai_service.generate_cover_letter(
                cv_data=cv_data,
                job_description=job_description,
                company_name=company_name,
                job_title=job_title,
                user_preferences=user_preferences
            )
            
            return {
                "success": True,
                "cover_letter": cover_letter_result.get("cover_letter", ""),
                "metadata": cover_letter_result.get("metadata", {}),
                "status": cover_letter_result.get("status", "completed")
            }
            
        except Exception as e:
            logger.error(f"Error generating cover letter: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to generate cover letter: {str(e)}"
            )

    async def get_user_documents(self, user_id: str, document_type: str = None) -> List[Dict[str, Any]]:
        """
        Get all documents for a user
        """
        try:
            db = await get_database()
            
            # Build query
            query = {
                "user_id": user_id,
                "is_active": True
            }
            if document_type:
                query["document_type"] = document_type
            
            # Execute query
            cursor = db.documents.find(
                query,
                {"text_content": 0}  # Exclude large text content
            ).sort("created_at", -1)
            
            documents = []
            async for doc in cursor:
                doc["_id"] = str(doc["_id"])
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            logger.error(f"Error fetching user documents: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to fetch documents: {str(e)}"
            )

    async def get_document(self, document_id: str, user_id: str) -> Optional[Dict[str, Any]]:
        """
        Get specific document
        """
        try:
            db = await get_database()
            
            doc = await db.documents.find_one({
                "_id": self._to_object_id(document_id),
                "user_id": user_id,
                "is_active": True
            })
            
            if doc:
                doc["_id"] = str(doc["_id"])
                return doc
            
            return None
            
        except Exception as e:
            logger.error(f"Error fetching document: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to fetch document: {str(e)}"
            )

    async def delete_document(self, document_id: str, user_id: str) -> bool:
        """
        Delete document (soft delete)
        """
        try:
            db = await get_database()
            
            result = await db.documents.update_one(
                {
                    "_id": self._to_object_id(document_id),
                    "user_id": user_id
                },
                {"$set": {"is_active": False, "updated_at": datetime.utcnow()}}
            )
            
            return result.modified_count > 0
            
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to delete document: {str(e)}"
            )

    def _validate_file(self, file: UploadFile) -> None:
        """
        Validate uploaded file
        """
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in self.allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"File type {file_ext} not supported. Allowed: {', '.join(self.allowed_extensions)}"
            )
        
        # Check file size (Note: file.size might not be available in all cases)
        # We'll check size after reading the content

    async def _save_file(self, file_content: bytes, filename: str, user_id: str) -> Path:
        """
        Save file content to disk
        """
        # Validate file size
        if len(file_content) > self.max_file_size:
            raise HTTPException(
                status_code=413,
                detail=f"File size exceeds {self.max_file_size / (1024*1024):.1f}MB limit"
            )
        
        # Generate unique filename
        file_ext = Path(filename).suffix
        unique_filename = f"{user_id}_{uuid.uuid4()}{file_ext}"
        file_path = self.upload_dir / unique_filename
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        return file_path

    def _extract_text(self, file_content: bytes, file_extension: str) -> str:
        """
        Extract text from different file types
        """
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_content)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_docx_text(file_content)
            elif file_extension == '.txt':
                return self._extract_txt_text(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text from {file_extension}: {str(e)}"
            )

    def _extract_pdf_text(self, file_content: bytes) -> str:
        """
        Extract text from PDF
        """
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text_parts = []
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            
            full_text = '\n'.join(text_parts)
            
            if not full_text.strip():
                raise ValueError("No readable text found in PDF")
            
            return full_text.strip()
            
        except Exception as e:
            raise ValueError(f"PDF processing error: {str(e)}")

    def _extract_docx_text(self, file_content: bytes) -> str:
        """
        Extract text from DOCX
        """
        try:
            doc = docx.Document(BytesIO(file_content))
            text_parts = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_parts)
            
            if not full_text.strip():
                raise ValueError("No readable text found in document")
            
            return full_text.strip()
            
        except Exception as e:
            raise ValueError(f"DOCX processing error: {str(e)}")

    def _extract_txt_text(self, file_content: bytes) -> str:
        """
        Extract text from TXT file
        """
        try:
            # Try UTF-8 first
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                text = file_content.decode('latin-1')
            except UnicodeDecodeError:
                raise ValueError("Unable to decode text file")
        
        if not text.strip():
            raise ValueError("Text file is empty")
        
        return text.strip()

    def _to_object_id(self, id_str: str):
        """
        Convert string ID to MongoDB ObjectId
        """
        from bson import ObjectId
        try:
            return ObjectId(id_str)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid document ID format")


# Singleton instance
document_service = DocumentService()